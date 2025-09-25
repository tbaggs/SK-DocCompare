import io
import re
from typing import Optional

import httpx
from pypdf import PdfReader
from docx import Document


def _is_url(path: str) -> bool:
    return bool(re.match(r"^https?://", path, re.IGNORECASE))


def _clean_text(text: str) -> str:
    # Normalize whitespace; keep it readable for prompts
    return re.sub(r"\s+", " ", text).strip()


async def load_text_from_source(path_or_url: str) -> str:
    """
    Load text from:
    - Azure Blob SAS URL (http(s))
    - Local file path
    Supports .txt, .pdf, .docx
    """
    if _is_url(path_or_url):
        async with httpx.AsyncClient(timeout=60.0) as client:
            resp = await client.get(path_or_url)
            resp.raise_for_status()
            data = resp.content
            # Guess by URL suffix; fall back to text
            lower = path_or_url.lower()
            if lower.endswith(".pdf"):
                reader = PdfReader(io.BytesIO(data))
                text = "\n".join(page.extract_text() or "" for page in reader.pages)
                return _clean_text(text)
            if lower.endswith(".docx"):
                doc = Document(io.BytesIO(data))
                text = "\n".join(p.text for p in doc.paragraphs)
                return _clean_text(text)
            # default treat as text
            return _clean_text(data.decode("utf-8", errors="ignore"))
    else:
        lower = path_or_url.lower()
        if lower.endswith(".pdf"):
            with open(path_or_url, "rb") as f:
                reader = PdfReader(f)
                text = "\n".join(page.extract_text() or "" for page in reader.pages)
                return _clean_text(text)
        if lower.endswith(".docx"):
            doc = Document(path_or_url)
            text = "\n".join(p.text for p in doc.paragraphs)
            return _clean_text(text)
        # default: text file
        with open(path_or_url, "r", encoding="utf-8", errors="ignore") as f:
            return _clean_text(f.read())
